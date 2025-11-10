# DEPENDENCIES
import io
import re
import PyPDF2
from typing import Any
from typing import Dict
from typing import Union
from pathlib import Path
from docx import Document
from typing import Optional


try:
    # For PyMuPDF
    import fitz  
    PYMUPDF_AVAILABLE = True

except ImportError:
    PYMUPDF_AVAILABLE = False
    print("[DocumentReader] PyMuPDF not available. Install with: pip install PyMuPDF")

# Encoding detection
try:
    import chardet
    CHARDET_AVAILABLE = True

except ImportError:
    CHARDET_AVAILABLE = False



class DocumentReader:
    """
    Document reader supporting PDF and DOCX : Uses PyMuPDF for better PDF extraction when available
    """
    # File Size Constraint : 10MB
    MAX_FILE_SIZE   = 10 * 1024 * 1024  

    # File Type Constraint
    ALLOWED_TYPES   = ["pdf", "docx", "doc"]

    # Minimum extracted text length
    MIN_TEXT_LENGTH = 100  
    

    @staticmethod
    def read_file(file_path_or_bytes: Union[str, Path, io.BytesIO], file_type: str = "pdf") -> str:
        """
        Read document and extract text with validation
        
        Arguments:
        ----------
            file_path_or_bytes { str / Path / BytesIO } : File path (str/Path) or bytes object (io.BytesIO)
            
            file_type                   { str }         : "pdf" or "docx"
        
        Returns:
        --------
                            { str }                     : Extracted and cleaned text
            
        Raises:
        -------
            ValueError                                  : If file type unsupported or validation fails

            Exception                                   : If extraction fails
        """
        # Validate file type
        if file_type.lower() not in DocumentReader.ALLOWED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}. Allowed types: {', '.join(DocumentReader.ALLOWED_TYPES)}")
        
        # Validate file size
        DocumentReader._validate_file_size(file_path_or_bytes = file_path_or_bytes)
        
        # Route to appropriate reader
        if (file_type.lower() == "pdf"):
            text = DocumentReader._read_pdf(file_or_bytes = file_path_or_bytes)

        elif (file_type.lower() in ["docx", "doc"]):
            text = DocumentReader._read_docx(file_or_bytes = file_path_or_bytes)
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Validate extracted text
        if (len(text.strip()) < DocumentReader.MIN_TEXT_LENGTH):
            raise ValueError(f"Extracted text too short ({len(text)} chars). Minimum: {DocumentReader.MIN_TEXT_LENGTH} chars. File may be corrupted or empty.")
        
        return text

    
    @staticmethod
    def _validate_file_size(file_path_or_bytes: Union[str, Path, io.BytesIO]) -> None:
        """
        Validate file size is within limits
        """
        if isinstance(file_path_or_bytes, (str, Path)):
            size = Path(file_path_or_bytes).stat().st_size
        
        else:
            # It's a file-like object
            current_pos = file_path_or_bytes.tell()
            # Seek to end
            file_path_or_bytes.seek(0, 2)  
            
            size        = file_path_or_bytes.tell()

            # Reset to original position
            file_path_or_bytes.seek(current_pos)  
        
        if (size > DocumentReader.MAX_FILE_SIZE):
            raise ValueError(f"File too large: {size / (1024 * 1024):.1f}MB. Maximum allowed: {DocumentReader.MAX_FILE_SIZE / (1024 * 1024):.1f}MB")
        
        if (size == 0):
            raise ValueError("File is empty (0 bytes)")
    

    @staticmethod
    def _read_pdf(file_or_bytes: Union[str, Path, io.BytesIO]) -> str:
        """
        Read PDF with PyMuPDF (preferred) or PyPDF2 (fallback)
        """
        if PYMUPDF_AVAILABLE:
            try:
                return DocumentReader._read_pdf_pymupdf(file_or_bytes = file_or_bytes)
            
            except Exception as e:
                print(f"[DocumentReader] PyMuPDF failed: {e}, falling back to PyPDF2")
                return DocumentReader._read_pdf_pypdf2(file_or_bytes)
        
        else:
            return DocumentReader._read_pdf_pypdf2(file_or_bytes = file_or_bytes)
    

    @staticmethod
    def _read_pdf_pymupdf(file_or_bytes: Union[str, Path, io.BytesIO]) -> str:
        """
        Read PDF using PyMuPDF (superior text extraction)
        """
        # Handle both file paths and bytes
        if isinstance(file_or_bytes, (str, Path)):
            doc = fitz.open(file_or_bytes)
        
        else:
            # It's a file-like object
            file_or_bytes.seek(0)
            file_content = file_or_bytes.read()
            doc          = fitz.open(stream   = file_content, 
                                     filetype = "pdf",
                                    )
        
        text = ""

        for page_num in range(doc.page_count):
            page      = doc[page_num]
            
            # Extract text with layout preservation
            page_text = page.get_text("text", sort = True)
            
            # Clean the text
            page_text = DocumentReader._clean_extracted_text(text = page_text)
            text     += page_text + "\n\n"
        
        doc.close()
        
        # Post-process entire document
        text = DocumentReader._post_process_text(text = text)
        
        return text
    

    @staticmethod
    def _read_pdf_pypdf2(file_or_bytes: Union[str, Path, io.BytesIO]) -> str:
        """
        Read PDF using PyPDF2 (fallback)
        """
        try:
            # Handle both file paths and bytes
            if isinstance(file_or_bytes, (str, Path)):
                with open(file_or_bytes, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text       = DocumentReader._extract_from_pypdf2(pdf_reader = pdf_reader)
            
            else:
                file_or_bytes.seek(0)
                pdf_reader = PyPDF2.PdfReader(stream = file_or_bytes)
                text       = DocumentReader._extract_from_pypdf2(pdf_reader = pdf_reader)
            
            text = DocumentReader._post_process_text(text = text)
            
            return text
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {repr(e)}")
    

    @staticmethod
    def _extract_from_pypdf2(pdf_reader: PyPDF2.PdfReader) -> str:
        """
        Extract text from PyPDF2 reader
        """
        text = ""

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            
            if page_text:
                page_text = DocumentReader._clean_extracted_text(text = page_text)
                text     += page_text + "\n\n"

        return text
    

    @staticmethod
    def _read_docx(file_or_bytes: Union[str, Path, io.BytesIO]) -> str:
        """
        Read DOCX file
        """
        try:
            # Handle both file paths and bytes
            if isinstance(file_or_bytes, (str, Path)):
                doc = Document(file_or_bytes)
            
            else:
                file_or_bytes.seek(0)
                doc = Document(file_or_bytes)
            
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    clean_text = DocumentReader._clean_extracted_text(text = paragraph.text)
                    text      += clean_text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            clean_text = DocumentReader._clean_extracted_text(text = cell.text)
                            text      += clean_text + " "

                    text += "\n"
            
            text = DocumentReader._post_process_text(text = text)
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {repr(e)}")
    

    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """
        Clean and normalize extracted text
        """
        if not text:
            return ""
        
        # Replace multiple newlines with single newline
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Fix hyphenated words split across lines
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        # Normalize whitespace (but preserve single newlines)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page numbers on separate lines
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        
        # Remove lines with just numbers
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        return text.strip()
    

    @staticmethod
    def _post_process_text(text: str) -> str:
        """
        Post-process entire extracted text
        """
        if not text:
            return ""
        
        # Remove excessive empty lines
        text  = re.sub(r'\n{3,}', '\n\n', text)
        
        # Fix mid-sentence line breaks (lowercase to lowercase)
        text  = re.sub(r'(?<=[a-z,])\n(?=[a-z])', ' ', text)
        
        # Ensure proper spacing around section numbers
        text  = re.sub(r'(\d+\.\d+)([A-Za-z])', r'\1 \2', text)
        
        # Remove excessive spaces
        text  = re.sub(r' {2,}', ' ', text)
        
        # Remove whitespace at line starts/ends
        lines = [line.strip() for line in text.split('\n')]
        text  = '\n'.join(lines)
        
        return text.strip()
    

    @staticmethod
    def extract_metadata(file_path_or_bytes: Union[str, Path, io.BytesIO], file_type: str = "pdf") -> Dict[str, Any]:
        """
        Extract document metadata (pages, author, creation date, etc.)
        
        Arguments:
        ----------
            file_path_or_bytes { str / Path / BytesIO } : File path or bytes object
            
            file_type                  { str }          : "pdf" or "docx"
        
        Returns:
        --------
                        { dict }                        : Dictionary containing metadata
        """
        metadata = {"pages"     : 0,
                    "title"     : "",
                    "author"    : "",
                    "creator"   : "",
                    "created"   : "",
                    "modified"  : "",
                    "file_type" : file_type,
                   }
        
        try:
            if ((file_type == "pdf") and PYMUPDF_AVAILABLE):
                if isinstance(file_path_or_bytes, (str, Path)):
                    doc = fitz.open(file_path_or_bytes)
                
                else:
                    file_path_or_bytes.seek(0)
                    file_content = file_path_or_bytes.read()
                    doc          = fitz.open(stream=file_content, filetype="pdf")
                
                metadata.update({"pages"    : doc.page_count,
                                 "title"    : doc.metadata.get("title", ""),
                                 "author"   : doc.metadata.get("author", ""),
                                 "creator"  : doc.metadata.get("creator", ""),
                                 "created"  : doc.metadata.get("creationDate", ""),
                                 "modified" : doc.metadata.get("modDate", ""),
                               })

                doc.close()
            
            elif (file_type in ["docx", "doc"]):
                if isinstance(file_path_or_bytes, (str, Path)):
                    doc = Document(file_path_or_bytes)
                
                else:
                    file_path_or_bytes.seek(0)
                    doc = Document(file_path_or_bytes)
                
                core_props = doc.core_properties
                metadata.update({"pages"    : len(doc.sections),
                                 "title"    : core_props.title or "",
                                 "author"   : core_props.author or "",
                                 "creator"  : core_props.author or "",
                                 "created"  : str(core_props.created) if core_props.created else "",
                                 "modified" : str(core_props.modified) if core_props.modified else "",
                               })
        
        except Exception as e:
            print(f"[DocumentReader] Metadata extraction failed: {repr(e)}")
        
        return metadata

    
    @staticmethod
    def detect_encoding(file_bytes: bytes) -> str:
        """
        Detect text encoding for better extraction
        
        Arguments:
        ----------
            file_bytes { bytes } : Raw file bytes
        
        Returns:
        --------
                { str }          : Detected encoding (e.g., 'utf-8', 'latin-1')
        """
        if not CHARDET_AVAILABLE:
            return 'utf-8'
        
        try:
            # Check first 10KB
            result = chardet.detect(file_bytes[:10000])  
            return result['encoding'] or 'utf-8'
        
        except Exception:
            return 'utf-8'
    

    @staticmethod
    def validate_file_integrity(file_path: Union[str, Path]) -> tuple[bool, str]:
        """
        Validate file isn't corrupted and is readable
        
        Arguments:
        ----------
            file_path { str } : Path to file
        
        Returns:
        --------
               { tuple }      : (is_valid, message) tuple
        """
        try:
            file_path = Path(file_path)
            
            # Check file exists
            if not file_path.exists():
                return False, "File does not exist"
            
            # Check file size
            file_size = file_path.stat().st_size
            
            if (file_size == 0):
                return (False, "File is empty (0 bytes)")
            
            # Less than 1KB
            if (file_size < 1024):  
                return (False, f"File suspiciously small ({file_size} bytes)")
            
            # Check file is readable
            with open(file_path, 'rb') as f:
                # Try reading first KB
                f.read(1024)  
            
            return (True, "File integrity OK")
            
        except PermissionError:
            return (False, "Permission denied - cannot read file")
        
        except Exception as e:
            return (False, f"File integrity check failed: {repr(e)}")